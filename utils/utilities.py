import streamlit as st

import pytube
from pytube import YouTube
import re
import os
import io
from streamlit.runtime.scriptrunner import add_script_run_ctx


import transformers
from transformers import pipeline

import soundfile as sf
import ffmpeg
import time
import tempfile
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd

import whisper
import transformers
from transformers import pipeline

import torch
import torchaudio
from torchaudio.transforms import Resample

from typing import Tuple, Union, List

from fpdf import FPDF

from docx import Document
from docx.shared import Pt

def plot_soundwave(df: pd.DataFrame, figsize: Tuple[int, int] = (3, 1)) -> plt.Figure:
    """
    Plots the soundwave from a DataFrame.

    Args:
    - df (pd.DataFrame): DataFrame containing 'Seconds' and 'Amplitude' columns.
    - figsize (Tuple[int, int]): The size of the figure in inches. Defaults to (3, 1).

    Returns:
    - fig (plt.Figure): Matplotlib figure object containing the plot.
    """
    fig, ax = plt.subplots(figsize=figsize, facecolor="None")
    ax.set_facecolor("None")
    ax.plot(df['Seconds'], df['Amplitude'], color='#34ddd7')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.get_xaxis().set_ticks([])
    ax.get_yaxis().set_ticks([])

    return fig


@st.cache_resource
def load_speech_transcriptor():
    return pipeline("automatic-speech-recognition")

def upsample_audio(
    frames, #: List[bytes],
    current_sample_rate: int = 48000,
    target_sampling_rate: int = 16000
) -> str:
    """
    Upsample audio from binary data to the target sampling rate using torchaudio.

    Parameters:
    - frames (List[bytes]): List of binary audio frames.
    - current_sample_rate (int): Current sample rate of the input audio (default: 48000).
    - target_sampling_rate (int): Target sampling rate for resampling (default: 16000).

    Returns:
    - output_path (str): Path to the upsampled audio file in WAV format.
    """
    # Save to a temporary WAV file
    temp_wav_file = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
    temp_wav_file_path = temp_wav_file.name
    temp_wav_file.close()

    # Convert frames to numpy array
    audio_data = np.frombuffer(b''.join(frames), dtype=np.int16)

    sf.write(temp_wav_file_path, audio_data, current_sample_rate)

    original_audio, _ = torchaudio.load(temp_wav_file_path, normalize=True)

    # Use torchaudio's Resample transform to upsample the audio
    resampler = Resample(current_sample_rate, target_sampling_rate)
    upsampled_audio = resampler(original_audio)

    output_path = os.path.join(tempfile.gettempdir(), "audio_upsampled.wav")
    torchaudio.save(output_path, upsampled_audio, sample_rate=target_sampling_rate)

    # Remove the temporary WAV file
    os.remove(temp_wav_file_path)
    return output_path

@st.cache_data
def create_plot_dataframe(_elapsed_time: torch.Tensor, _accumulated_waveform: torch.Tensor, idx: int) -> pd.DataFrame:
    """
    Create a DataFrame with columns "Seconds" and "Amplitude" for plotting.

    Args:
        _elapsed_time (torch.Tensor): A tensor containing the elapsed time.
        _accumulated_waveform (torch.Tensor): A tensor containing the accumulated waveform.
        idx (int): Index value.

    Returns:
        pd.DataFrame: A DataFrame containing the elapsed time and amplitude for plotting.
    """
    df = pd.DataFrame({
        "Seconds": _elapsed_time.numpy(),
        "Amplitude": _accumulated_waveform.t().numpy().flatten()
    })
    return df

def start_recording():
    """
    Starts the audio recording process by setting the 'run' flag to True in the session state.
    """
    st.session_state['run'] = True

def stop_recording():
    """
    Stops the audio recording process by setting the 'run' flag to False in the session state.
    Stores the current transcript in 'on_screen_transcript' to make it remanent and clears the 'transcript' variable.
    """
    st.session_state['run'] = False
    st.session_state['on_screen_transcript'] = st.session_state['transcript']
    st.session_state['transcript'] = ""

@st.cache_resource
def load_speech_to_text_model(device: Union[str, torch.device] = "cpu"):
    """
    Load the whisper model for speech-to-text.

    Parameters:
        device (str): Device to use for the model, e.g., "cpu" or "cuda".

    Returns:
        torch.nn.Module: Loaded whisper model.
    """
    if device == "cuda" and torch.cuda.is_available():
        
        model = whisper.load_model("tiny", device="cuda")

    else:
        model = whisper.load_model("tiny", device="cpu")

    return model


def fetch_audio_from_YT(YT_Video_URL: str) -> Union[str, Tuple[str, str]]:
  """
    Retrieve the audio file and video title from a YouTube video link.

    Args:
        YT_Video_URL (str): The link of the YouTube video from which the audio file shall be retrieved.

    Returns:
        Union[None, Tuple[str, str]]: Returns either None if it did not manage to fetch the audio file,
                                        or a tuple containing the video title and the path to the downloaded audio file.
  Note: Issue with pytube (https://stackoverflow.com/questions/68680322/pytube-urllib-error-httperror-http-error-410-gone)
  """

  try:
    video = YouTube(YT_Video_URL)

    # Obtain the mp4 audio stream of the desired video
    audio_stream = video.streams.filter(only_audio=True, file_extension='mp4').first().download()
    audio_name = video.title
    return audio_stream, audio_name

  except:
    # print("Error: failed to fecth")
    return None, "Error: failed to fecth video"
  
@st.cache_data
def generate_video_transcript(audio, _model, device) -> str:
  """
  Genrate the transcript from a YouTube video, based on its fetched audio file.

  Arg: audio (pytube.streams.Stream): The audio file of the YouTube video that you want the transcript.
  Return (str): Returns  the trasncribed video to text format, as a string.
  Note: https://stackoverflow.com/questions/73845566/openai-whisper-filenotfounderror-winerror-2-the-system-cannot-find-the-file
  """
  if device =="cuda" and torch.cuda.is_available():
    # Use FP16 if on GPU
    _model = _model.to("cuda")
    transcript = _model.transcribe(audio)

  else: 
    # Disable FP16 if on CPU
    _model = _model.to("cpu")
    transcript = _model.transcribe(audio, fp16=False)

  text_transcript = transcript["text"]

  return text_transcript


def extract_file_name(file_path: str) -> str:
    """
    Extracts the file name from a given file path.

    Args:
        file_path (str): The file path from which to extract the file name.

    Returns:
        str: The extracted file name without the file extension, or None if the file path is invalid.
    """
    match = re.match(r"^(.+?)(\.[^.]+)?$", os.path.basename(file_path))
    if match:
        return match.group(1)
    else:
        return None


def click_button(button):
    # Reset all buttons to False
    for key in st.session_state.clicked.keys():
        st.session_state.clicked[key] = False

    # Set the clicked button to True
    st.session_state.clicked[button] = True

@st.cache_data
def generate_pdf(
        transcript: str,
        audio_file: str,
        font: str = "Arial"
) -> FPDF:
    """
    Generate a PDF document with a title centered at the top and the transcript below it.

    Parameters:
    - transcript (str): The transcript content to be included in the PDF.
    - audio_file (str): The name of the audio file for which the transcript is generated.
    - font (str, optional): The font style to be used (default is "Arial").

    Returns:
    FPDF: An instance of the FPDF class representing the generated PDF document.
    """
    # Create PDF object
    pdf = FPDF()

    # Add a page
    pdf.add_page()

    # Set title font size and style
    title_font_size = 15
    pdf.set_font(font, size=title_font_size)
    
    # Create the title
    title = "Transcript of the " + audio_file + " audio file"
    pdf.cell(200, 10, txt=title, ln=1, align='C')

    # Set transcript font size and style
    transcript_font_size = 11
    pdf.set_font(font, size=transcript_font_size)

    # Add the transcript
    pdf.multi_cell(0, 10, txt=transcript)

    return pdf


def generate_word_transcript(transcript: str, audio_file: str, font: str = "Arial") -> Document:
    """
    Generate a Word document with a title centered at the top and the transcript below it.

    Parameters:
    - transcript (str): The transcript content to be included in the Word document.
    - audio_file (str): The name of the audio file for which the transcript is generated.
    - font (str, optional): The font style to be used (default is "Calibri").

    Returns:
    Document: An instance of the Document class representing the generated Word document.
    """
    # Create Word document object
    doc = Document()

    # Set title font size and style
    title_font_size = 15
    
    # Create the title
    title = "Transcript of the " + audio_file + " audio file"
    title_paragraph = doc.add_paragraph(title)
    title_runs = title_paragraph.runs
    title_run = title_runs[0]  # Assuming there is only one run in the paragraph
    title_run.font.size = Pt(title_font_size)  # Pt is used to specify font size in points

    # Set transcript font size and style
    transcript_font_size = 11
    
    # Add the transcript
    transcript_paragraph = doc.add_paragraph(transcript)
    transcript_runs = transcript_paragraph.runs
    for run in transcript_runs:
        run.font.size = Pt(transcript_font_size)  # Pt is used to specify font size in points

    return doc